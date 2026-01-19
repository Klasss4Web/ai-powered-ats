import os
import json
import sqlite3
import hashlib
import jwt
import datetime
from functools import wraps
from flask import Flask, request, jsonify, send_file, g
from flask_cors import CORS
import PyPDF2
from dotenv import load_dotenv
# Install required library: pip install python-docx
from docx import Document
from io import BytesIO
import requests

load_dotenv()
# 1. Gemini API Setup
# Requires: pip install google-genai
try:
    import google.generativeai as genai
    api_key = os.getenv("GEMINI_API_KEY")
    if api_key:
        genai.configure(api_key=api_key)
    else:
        print("WARNING: GEMINI_API_KEY not set in environment.")
    try:
        model = genai.GenerativeModel('gemini-robotics-er-1.5-preview')
        print("Gemini GenerativeModel initialized successfully.")
    except Exception:
        model = None
except Exception as e:
    # If API key is missing or invalid or import fails
    print(f"ERROR: Could not initialize Gemini Client. Check if GEMINI_API_KEY is set. Details: {e}")
    model = None

# 2. Database Setup
DATABASE = 'ats_matcher.db'
JWT_SECRET_KEY = os.getenv('JWT_SECRET_KEY', 'your-secret-key-change-in-production')

def get_db():
    """Get database connection."""
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = sqlite3.connect(DATABASE)
        db.row_factory = sqlite3.Row
    return db

def init_db():
    """Initialize database tables."""
    with app.app_context():
        db = get_db()
        cursor = db.cursor()

        # Create users table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                name TEXT NOT NULL,
                subscription_type TEXT DEFAULT 'free' CHECK (subscription_type IN ('free', 'premium')),
                subscription_expires_at TIMESTAMP NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        # Add reset_token and reset_expires columns if they don't exist
        try:
            cursor.execute('ALTER TABLE users ADD COLUMN reset_token TEXT')
        except:
            pass
        try:
            cursor.execute('ALTER TABLE users ADD COLUMN reset_expires TIMESTAMP')
        except:
            pass

        # Create usage tracking table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS usage_tracking (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                action_type TEXT NOT NULL,
                date_created DATE DEFAULT CURRENT_DATE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                metadata TEXT,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')

        # Add metadata column if it doesn't exist (for existing databases)
        try:
            cursor.execute('ALTER TABLE usage_tracking ADD COLUMN metadata TEXT')
        except:
            pass  # Column might already exist

        # Create sessions table for refresh tokens (optional)
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                token TEXT UNIQUE NOT NULL,
                expires_at TIMESTAMP NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')

        db.commit()

# 3. Authentication Functions
def hash_password(password):
    """Hash password using SHA-256."""
    return hashlib.sha256(password.encode()).hexdigest()

def generate_token(user_id, email):
    """Generate JWT token."""
    payload = {
        'user_id': user_id,
        'email': email,
        'exp': datetime.datetime.utcnow() + datetime.timedelta(hours=24),
        'iat': datetime.datetime.utcnow()
    }
    return jwt.encode(payload, JWT_SECRET_KEY, algorithm='HS256')

def verify_token(token):
    """Verify JWT token."""
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=['HS256'])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

def token_required(f):
    """Decorator to require authentication token."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = None

        # Get token from header
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            if auth_header.startswith('Bearer '):
                token = auth_header.split(' ')[1]

        if not token:
            return jsonify({'error': 'Token is missing'}), 401

        payload = verify_token(token)
        if not payload:
            return jsonify({'error': 'Token is invalid or expired'}), 401

        # Add user info to request
        g.user_id = payload['user_id']
        g.user_email = payload['email']

        return f(*args, **kwargs)
    return decorated_function

def check_usage_limit(user_id, action_type='analysis'):
    """Check if user has exceeded their daily usage limit."""
    db = get_db()
    cursor = db.cursor()

    # Get user's subscription info
    cursor.execute('SELECT subscription_type, subscription_expires_at FROM users WHERE id = ?', (user_id,))
    user = cursor.fetchone()

    if not user:
        return False, "User not found"

    subscription_type = user['subscription_type']
    expires_at = user['subscription_expires_at']

    # Check if subscription is expired
    if subscription_type == 'premium' and expires_at:
        try:
            expires_dt = datetime.datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
            if datetime.datetime.utcnow() > expires_dt:
                subscription_type = 'free'  # Treat as free if expired
        except:
            pass  # If date parsing fails, assume still valid

    # Get today's usage count for analyses
    today = datetime.date.today().isoformat()
    cursor.execute('''
        SELECT COUNT(*) as count FROM usage_tracking
        WHERE user_id = ? AND action_type = ? AND date_created = ?
    ''', (user_id, action_type, today))
    usage = cursor.fetchone()

    current_usage = usage['count'] if usage else 0

    # Check for pay-as-you-go payments today
    cursor.execute('''
        SELECT COUNT(*) as payment_count FROM usage_tracking
        WHERE user_id = ? AND action_type = ? AND date_created = ?
    ''', (user_id, 'payment', today))
    payments = cursor.fetchone()
    payment_count = payments['payment_count'] if payments else 0

    # Define limits
    limits = {
        'free': 1,
        'premium': 10
    }

    limit = limits.get(subscription_type, 1)

    # Allow additional analyses for each pay-as-you-go payment
    effective_limit = limit + payment_count

    if current_usage >= effective_limit:
        if subscription_type == 'free':
            return False, f"Daily {action_type} limit exceeded. Pay-as-you-go allows {payment_count} {action_type}(s) per day."
        else:
            return False, f"Daily {action_type} limit exceeded. {subscription_type.capitalize()} users can perform {limit} {action_type}(s) per day, plus {payment_count} pay-as-you-go."

    return True, current_usage

def record_usage(user_id, action_type='analysis'):
    """Record usage for tracking purposes."""
    db = get_db()
    cursor = db.cursor()

    today = datetime.date.today().isoformat()
    try:
        cursor.execute('''
            INSERT INTO usage_tracking (user_id, action_type, date_created)
            VALUES (?, ?, ?)
        ''', (user_id, action_type, today))
        db.commit()
    except sqlite3.IntegrityError:
        # If unique constraint fails, usage already recorded for today
        pass

# 4. Flask App Setup
# Requires: pip install flask flask-cors PyPDF2
app = Flask(__name__)
CORS(app) # Allow all origins for development purposes

@app.teardown_appcontext
def close_connection(exception):
    """Close database connection."""
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()


@app.route('/', methods=['GET'])
def index():
    """Basic landing page so visiting root doesn't return 404."""
    return (
        '<h2>ATS Matcher Backend</h2>'
        '<p>Available endpoints:</p>'
        '<ul>'
        '<li>POST <code>/api/match</code> - Upload resume (file) and job_description (form field)</li>'
        '<li>POST <code>/api/payment/initialize</code> - Initialize payment transaction</li>'
        '<li>GET <code>/api/payment/verify/&lt;reference&gt;</code> - Verify payment transaction</li>'
        '<li>GET <code>/api/payment/config</code> - Get payment configuration</li>'
        '<li>GET <code>/health</code> - Health check</li>'
        '</ul>'
    )


@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "ok"})

# Authentication Endpoints
@app.route('/api/auth/register', methods=['POST'])
def register():
    """Register a new user."""
    try:
        data = request.get_json()

        if not data or not data.get('email') or not data.get('password') or not data.get('name'):
            return jsonify({'error': 'Email, password, and name are required'}), 400

        email = data['email'].strip().lower()
        password = data['password']
        name = data['name'].strip()

        if len(password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters long'}), 400

        db = get_db()
        cursor = db.cursor()

        # Check if user already exists
        cursor.execute('SELECT id FROM users WHERE email = ?', (email,))
        if cursor.fetchone():
            return jsonify({'error': 'User with this email already exists'}), 409

        # Create new user
        password_hash = hash_password(password)
        cursor.execute(
            'INSERT INTO users (email, password_hash, name, subscription_type) VALUES (?, ?, ?, ?)',
            (email, password_hash, name, 'free')
        )
        user_id = cursor.lastrowid
        db.commit()

        # Generate token
        token = generate_token(user_id, email)

        return jsonify({
            'message': 'User registered successfully',
            'token': token,
            'user': {
                'id': user_id,
                'email': email,
                'name': name,
                'subscription_type': 'free',
                'subscription_expires_at': None
            }
        }), 201

    except Exception as e:
        print(f"Registration error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    """Login user."""
    try:
        data = request.get_json()

        if not data or not data.get('email') or not data.get('password'):
            return jsonify({'error': 'Email and password are required'}), 400

        email = data['email'].strip().lower()
        password = data['password']

        db = get_db()
        cursor = db.cursor()

        # Find user
        cursor.execute('SELECT id, email, password_hash, name, subscription_type, subscription_expires_at FROM users WHERE email = ?', (email,))
        user = cursor.fetchone()

        if not user or user['password_hash'] != hash_password(password):
            return jsonify({'error': 'Invalid email or password'}), 401

        # Generate token
        token = generate_token(user['id'], user['email'])

        return jsonify({
            'message': 'Login successful',
            'token': token,
            'user': {
                'id': user['id'],
                'email': user['email'],
                'name': user['name'],
                'subscription_type': user['subscription_type'],
                'subscription_expires_at': user['subscription_expires_at']
            }
        }), 200

    except Exception as e:
        print(f"Login error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/auth/verify', methods=['GET'])
@token_required
def verify_token_endpoint():
    """Verify if token is valid."""
    return jsonify({
        'valid': True,
        'user': {
            'id': g.user_id,
            'email': g.user_email
        }
    }), 200

@app.route('/api/auth/logout', methods=['POST'])
@token_required
def logout():
    """Logout user (client-side token removal is main logout mechanism)."""
    # In a more complex system, you might want to blacklist tokens
    return jsonify({'message': 'Logged out successfully'}), 200

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    """Send password reset email."""
    try:
        data = request.get_json()
        if not data or not data.get('email'):
            return jsonify({'error': 'Email is required'}), 400

        email = data['email'].strip().lower()

        db = get_db()
        cursor = db.cursor()

        # Find user
        cursor.execute('SELECT id FROM users WHERE email = ?', (email,))
        user = cursor.fetchone()

        if not user:
            # Don't reveal if email exists or not for security
            return jsonify({'message': 'If the email exists, a reset link has been sent.'}), 200

        # Generate reset token
        import secrets
        reset_token = secrets.token_urlsafe(32)
        reset_expires = datetime.datetime.utcnow() + datetime.timedelta(hours=1)

        # Update user with reset token
        cursor.execute('''
            UPDATE users
            SET reset_token = ?, reset_expires = ?
            WHERE id = ?
        ''', (reset_token, reset_expires.isoformat(), user['id']))
        db.commit()

        # In a real app, send email here
        print(f"Password reset token for {email}: {reset_token}")
        print(f"Reset link: http://localhost:5173/reset-password?token={reset_token}")

        return jsonify({'message': 'If the email exists, a reset link has been sent.'}), 200

    except Exception as e:
        print(f"Forgot password error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/auth/reset-password', methods=['POST'])
def reset_password():
    """Reset password using token."""
    try:
        data = request.get_json()
        if not data or not data.get('token') or not data.get('new_password'):
            return jsonify({'error': 'Token and new password are required'}), 400

        token = data['token']
        new_password = data['new_password']

        if len(new_password) < 6:
            return jsonify({'error': 'Password must be at least 6 characters long'}), 400

        db = get_db()
        cursor = db.cursor()

        # Find user with valid token
        cursor.execute('''
            SELECT id FROM users
            WHERE reset_token = ? AND reset_expires > ?
        ''', (token, datetime.datetime.utcnow().isoformat()))
        user = cursor.fetchone()

        if not user:
            return jsonify({'error': 'Invalid or expired token'}), 400

        # Update password and clear reset token
        new_password_hash = hash_password(new_password)
        cursor.execute('''
            UPDATE users
            SET password_hash = ?, reset_token = NULL, reset_expires = NULL
            WHERE id = ?
        ''', (new_password_hash, user['id']))
        db.commit()

        return jsonify({'message': 'Password reset successfully'}), 200

    except Exception as e:
        print(f"Reset password error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/user/usage', methods=['GET'])
@token_required
def get_user_usage():
    """Get user's current usage status and limits."""
    try:
        db = get_db()
        cursor = db.cursor()

        # Get user's subscription info
        cursor.execute('SELECT subscription_type, subscription_expires_at FROM users WHERE id = ?', (g.user_id,))
        user = cursor.fetchone()

        if not user:
            return jsonify({'error': 'User not found'}), 404

        subscription_type = user['subscription_type']
        expires_at = user['subscription_expires_at']

        # Check if subscription is expired
        is_expired = False
        if subscription_type == 'premium' and expires_at:
            try:
                expires_dt = datetime.datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
                if datetime.datetime.utcnow() > expires_dt:
                    is_expired = True
                    subscription_type = 'free'  # Treat as free if expired
            except:
                pass

        # Get today's usage count
        today = datetime.date.today().isoformat()
        cursor.execute('''
            SELECT COUNT(*) as count FROM usage_tracking
            WHERE user_id = ? AND action_type = ? AND date_created = ?
        ''', (g.user_id, 'analysis', today))
        usage = cursor.fetchone()

        current_usage = usage['count'] if usage else 0

        # Check for pay-as-you-go payments today
        cursor.execute('''
            SELECT COUNT(*) as payment_count FROM usage_tracking
            WHERE user_id = ? AND action_type = ? AND date_created = ?
        ''', (g.user_id, 'payment', today))
        payments = cursor.fetchone()
        payment_count = payments['payment_count'] if payments else 0

        # Define limits
        limits = {
            'free': 1,
            'premium': 10
        }

        base_limit = limits.get(subscription_type, 1)
        effective_limit = base_limit + payment_count
        remaining = max(0, effective_limit - current_usage)

        return jsonify({
            'subscription_type': subscription_type,
            'subscription_expires_at': expires_at,
            'is_expired': is_expired,
            'current_usage': current_usage,
            'daily_limit': base_limit,
            'pay_as_you_go_payments': payment_count,
            'effective_limit': effective_limit,
            'remaining_analyses': remaining,
            'can_perform_analysis': remaining > 0
        }), 200

    except Exception as e:
        print(f"Usage check error: {e}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/pay-as-you-go', methods=['POST'])
@token_required
def pay_as_you_go():
    """Process pay-as-you-go payment for one analysis."""
    try:
        data = request.get_json()
        amount = data.get('amount', 1.00)

        # In a real implementation, this would integrate with a payment processor like Stripe
        # For now, we'll simulate a successful payment

        # Log the payment (in production, this would be handled by payment processor)
        db = get_db()
        cursor = db.cursor()

        # Check if payment already recorded (prevent duplicates)
        cursor.execute('''
            SELECT id FROM usage_tracking 
            WHERE user_id = ? AND action_type = ? AND metadata LIKE ? AND date_created = ?
        ''', (g.user_id, 'payment', f'%"type":"pay_as_you_go"%', datetime.date.today().isoformat()))
        existing = cursor.fetchone()
        
        if not existing:
            cursor.execute('''
                INSERT INTO usage_tracking (user_id, action_type, date_created, metadata)
                VALUES (?, ?, ?, ?)
            ''', (g.user_id, 'payment', datetime.date.today().isoformat(), json.dumps({
                'amount': amount,
                'type': 'pay_as_you_go'
            })))

            db.commit()
        else:
            print("Pay-as-you-go payment already recorded today, skipping duplicate")

        return jsonify({
            'success': True,
            'message': 'Payment processed successfully',
            'amount': amount
        }), 200

    except Exception as e:
        print(f"Pay-as-you-go error: {e}")
        return jsonify({'error': 'Payment processing failed'}), 500

# Payment Endpoints
PAYSTACK_SECRET_KEY = os.getenv('PAYSTACK_SECRET_KEY')
PAYSTACK_BASE_URL = 'https://api.paystack.co'
PAYSTACK_CALLBACK_URL = os.getenv('PAYSTACK_CALLBACK_URL', 'http://localhost:5173/')

# PayPal Configuration
PAYPAL_CLIENT_ID = os.getenv('PAYPAL_CLIENT_ID')
PAYPAL_CLIENT_SECRET = os.getenv('PAYPAL_CLIENT_SECRET')
PAYPAL_BASE_URL = 'https://api-m.sandbox.paypal.com'  # Use sandbox for testing

def get_paypal_access_token():
    """Get PayPal access token."""
    try:
        auth = (PAYPAL_CLIENT_ID, PAYPAL_CLIENT_SECRET)
        headers = {
            'Accept': 'application/json',
            'Accept-Language': 'en_US',
        }
        data = {'grant_type': 'client_credentials'}
        response = requests.post(f'{PAYPAL_BASE_URL}/v1/oauth2/token', auth=auth, headers=headers, data=data)
        response.raise_for_status()
        return response.json()['access_token']
    except Exception as e:
        print(f"PayPal auth error: {e}")
        return None

@app.route('/api/payment/initialize', methods=['POST'])
@token_required
def initialize_payment():
    """Initialize a payment transaction with Paystack or PayPal."""
    try:
        data = request.get_json()
        email = data.get('email')
        amount = data.get('amount')  # Amount in kobo for Paystack, cents for PayPal
        gateway = data.get('gateway', 'paystack')  # Default to paystack
        payment_type = data.get('payment_type', 'pay_as_you_go')  # 'pay_as_you_go' or 'subscription'
        plan_type = data.get('plan_type')  # 'monthly' or 'yearly' for subscriptions

        if not email or not amount:
            return jsonify({'error': 'Email and amount are required'}), 400

        if gateway == 'paypal':
            # PayPal integration
            if not PAYPAL_CLIENT_ID or not PAYPAL_CLIENT_SECRET:
                return jsonify({'error': 'PayPal payment service not configured'}), 500

            access_token = get_paypal_access_token()
            if not access_token:
                return jsonify({'error': 'Failed to authenticate with PayPal'}), 500

            # Convert amount to USD cents format for PayPal
            amount_usd = amount / 100  # Assuming amount comes in cents

            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {access_token}'
            }

            paypal_payload = {
                'intent': 'CAPTURE',
                'purchase_units': [{
                    'amount': {
                        'currency_code': 'USD',
                        'value': f'{amount_usd:.2f}'
                    },
                    'description': f'ATS Matcher {payment_type.replace("_", " ").title()}{" - " + plan_type if plan_type else ""}'
                }],
                'application_context': {
                    'return_url': PAYSTACK_CALLBACK_URL,
                    'cancel_url': f'{PAYSTACK_CALLBACK_URL}?payment=cancelled'
                }
            }

            response = requests.post(f'{PAYPAL_BASE_URL}/v2/checkout/orders', headers=headers, json=paypal_payload)
            result = response.json()

            if response.status_code == 201:
                # Find approval URL
                approval_url = None
                for link in result.get('links', []):
                    if link.get('rel') == 'approve':
                        approval_url = link['href']
                        break

                return jsonify({
                    'status': True,
                    'data': {
                        'authorization_url': approval_url,
                        'reference': result['id'],
                        'gateway': 'paypal'
                    }
                }), 200
            else:
                return jsonify({'error': 'PayPal payment initialization failed', 'details': result}), 400

        else:
            # Paystack integration (default)
            if not PAYSTACK_SECRET_KEY:
                return jsonify({'error': 'Paystack payment service not configured'}), 500

            url = f"{PAYSTACK_BASE_URL}/transaction/initialize"
            headers = {
                'Authorization': f'Bearer {PAYSTACK_SECRET_KEY}',
                'Content-Type': 'application/json'
            }
            payload = {
                'email': email,
                'amount': str(amount),
                'callback_url': PAYSTACK_CALLBACK_URL,
                'metadata': {
                    'user_id': g.user_id,
                    'type': payment_type,
                    'plan_type': plan_type,
                    'gateway': gateway
                }
            }

            response = requests.post(url, headers=headers, json=payload)
            result = response.json()

            if response.status_code == 200 and result.get('status'):
                return jsonify(result), 200
            else:
                return jsonify({'error': 'Payment initialization failed', 'details': result}), 400

    except Exception as e:
        print(f"Payment initialization error: {e}")
        return jsonify({'error': 'Payment initialization failed'}), 500

@app.route('/api/payment/verify/<reference>', methods=['GET'])
@token_required
def verify_payment(reference):
    """Verify a payment transaction with Paystack."""
    print(f"Paystack verification called for reference: {reference}, user: {g.user_id}")
    try:
        if not PAYSTACK_SECRET_KEY:
            return jsonify({'error': 'Payment service not configured'}), 500

        url = f"{PAYSTACK_BASE_URL}/transaction/verify/{reference}"
        headers = {
            'Authorization': f'Bearer {PAYSTACK_SECRET_KEY}'
        }

        response = requests.get(url, headers=headers)
        result = response.json()
        print(f"Paystack API response: {result}")

        if response.status_code == 200 and result.get('status') and result.get('data', {}).get('status') == 'success':
            # Get payment metadata
            metadata = result.get('data', {}).get('metadata', {})
            payment_type = metadata.get('type', 'pay_as_you_go')
            plan_type = metadata.get('plan_type')

            # Check if payment already processed
            db = get_db()
            cursor = db.cursor()
            cursor.execute('''
                SELECT id FROM usage_tracking 
                WHERE user_id = ? AND action_type = ? AND metadata LIKE ?
            ''', (g.user_id, 'payment', f'%"reference":"{reference}"%'))
            existing = cursor.fetchone()
            
            if not existing:
                if payment_type == 'subscription_upgrade':
                    # Upgrade user to premium
                    if plan_type == 'yearly':
                        expires_at = datetime.datetime.now() + datetime.timedelta(days=365)
                    else:  # monthly
                        expires_at = datetime.datetime.now() + datetime.timedelta(days=30)

                    cursor.execute('''
                        UPDATE users
                        SET subscription_type = 'premium', subscription_expires_at = ?
                        WHERE id = ?
                    ''', (expires_at.isoformat(), g.user_id))

                    db.commit()
                    print(f"User upgraded to premium ({plan_type})")
                else:
                    # Record pay-as-you-go payment
                    amount_naira = result.get('data', {}).get('amount')
                    cursor.execute('''
                        INSERT INTO usage_tracking (user_id, action_type, date_created, metadata)
                        VALUES (?, ?, ?, ?)
                    ''', (g.user_id, 'payment', datetime.date.today().isoformat(), json.dumps({
                        'amount': amount_naira,
                        'currency': 'NGN',
                        'type': payment_type,
                        'reference': reference,
                        'gateway': 'paystack'
                    })))

                    db.commit()
                    print("Payment recorded successfully")
            else:
                print("Payment already processed, skipping duplicate")
            
            return jsonify(result), 200
        else:
            print(f"Payment verification failed: {result}")
            return jsonify({'error': 'Payment verification failed', 'details': result}), 400

    except Exception as e:
        print(f"Payment verification error: {e}")
        return jsonify({'error': 'Payment verification failed'}), 500

@app.route('/api/payment/manual-verify/<reference>', methods=['POST'])
@token_required
def manual_verify_payment(reference):
    """Manually verify a payment transaction (for debugging)."""
    print(f"Manual verification called for reference: {reference}, user: {g.user_id}")
    try:
        data = request.get_json()
        gateway = data.get('gateway', 'paystack')

        if gateway == 'paypal':
            # For PayPal, the reference is the order ID
            return verify_paypal_payment(reference)
        else:
            # For Paystack
            return verify_payment(reference)

    except Exception as e:
        print(f"Manual verification error: {e}")
        return jsonify({'error': 'Manual verification failed'}), 500

@app.route('/api/payment/paypal/success', methods=['GET'])
def paypal_success():
    """Handle PayPal payment success redirect."""
    token = request.args.get('token')
    if token:
        # Redirect to frontend with PayPal order ID for verification
        return redirect(f'{PAYSTACK_CALLBACK_URL}?paypal_verify={token}')
    return redirect(PAYSTACK_CALLBACK_URL)

@app.route('/api/payment/paypal/cancel', methods=['GET'])
def paypal_cancel():
    """Handle PayPal payment cancellation."""
    return redirect(f'{PAYSTACK_CALLBACK_URL}?payment=cancelled')

@app.route('/api/payment/verify-paypal/<order_id>', methods=['GET'])
@token_required
def verify_paypal_payment(order_id):
    """Verify PayPal payment."""
    print(f"PayPal verification called for order_id: {order_id}, user: {g.user_id}")
    try:
        if not PAYPAL_CLIENT_ID or not PAYPAL_CLIENT_SECRET:
            return jsonify({'error': 'PayPal payment service not configured'}), 500

        access_token = get_paypal_access_token()
        if not access_token:
            return jsonify({'error': 'Failed to authenticate with PayPal'}), 500

        headers = {
            'Authorization': f'Bearer {access_token}'
        }

        # Get order details
        response = requests.get(f'{PAYPAL_BASE_URL}/v2/checkout/orders/{order_id}', headers=headers)
        result = response.json()
        print(f"PayPal order details: {result}")

        if response.status_code == 200 and result.get('status') == 'APPROVED':
            # Payment approved - now capture it
            capture_url = None
            for link in result.get('links', []):
                if link.get('rel') == 'capture':
                    capture_url = link['href']
                    break
            
            if capture_url:
                capture_response = requests.post(capture_url, headers=headers, json={})
                capture_result = capture_response.json()
                print(f"PayPal capture result: {capture_result}")
                
                if capture_response.status_code == 201 and capture_result.get('status') == 'COMPLETED':
                    # Get payment details
                    amount_usd = float(result.get('purchase_units', [{}])[0].get('amount', {}).get('value', 0))
                    description = result.get('purchase_units', [{}])[0].get('description', '')
                    
                    # Parse payment type from description
                    if 'subscription' in description.lower():
                        payment_type = 'subscription_upgrade'
                        plan_type = 'monthly' if 'monthly' in description.lower() else 'yearly'
                    else:
                        payment_type = 'pay_as_you_go'
                        plan_type = None

                    # Check if payment already processed
                    db = get_db()
                    cursor = db.cursor()
                    cursor.execute('''
                        SELECT id FROM usage_tracking 
                        WHERE user_id = ? AND action_type = ? AND metadata LIKE ?
                    ''', (g.user_id, 'payment', f'%"reference":"{order_id}"%'))
                    existing = cursor.fetchone()
                    
                    if not existing:
                        if payment_type == 'subscription_upgrade':
                            # Upgrade user to premium
                            if plan_type == 'yearly':
                                expires_at = datetime.datetime.now() + datetime.timedelta(days=365)
                            else:  # monthly
                                expires_at = datetime.datetime.now() + datetime.timedelta(days=30)

                            cursor.execute('''
                                UPDATE users
                                SET subscription_type = 'premium', subscription_expires_at = ?
                                WHERE id = ?
                            ''', (expires_at.isoformat(), g.user_id))

                            db.commit()
                            print(f"User upgraded to premium ({plan_type}) via PayPal")
                        else:
                            # Record pay-as-you-go payment
                            cursor.execute('''
                                INSERT INTO usage_tracking (user_id, action_type, date_created, metadata)
                                VALUES (?, ?, ?, ?)
                            ''', (g.user_id, 'payment', datetime.date.today().isoformat(), json.dumps({
                                'amount': amount_usd,
                                'currency': 'USD',
                                'type': payment_type,
                                'reference': order_id,
                                'gateway': 'paypal'
                            })))

                            db.commit()
                            print("PayPal payment recorded successfully")
                    else:
                        print("PayPal payment already processed, skipping duplicate")

                    return jsonify({
                        'status': True,
                        'data': {
                            'status': 'success',
                            'amount': amount_usd,
                            'currency': 'USD',
                            'reference': order_id
                        }
                    }), 200
                else:
                    print(f"PayPal capture failed: {capture_result}")
                    return jsonify({'error': 'PayPal payment capture failed', 'details': capture_result}), 400
            else:
                print("PayPal capture URL not found")
                return jsonify({'error': 'PayPal capture URL not found'}), 400
        else:
            print(f"PayPal verification failed: {result}")
            return jsonify({'error': 'PayPal payment verification failed', 'details': result}), 400

    except Exception as e:
        print(f"PayPal verification error: {e}")
        return jsonify({'error': 'PayPal payment verification failed'}), 500

@app.route('/api/payment/config', methods=['GET'])
def get_payment_config():
    """Get payment configuration for frontend."""
    paystack_public_key = os.getenv('PAYSTACK_PK_KEY')
    paypal_client_id = os.getenv('PAYPAL_CLIENT_ID')

    config = {}
    if paystack_public_key:
        config['paystack_public_key'] = paystack_public_key
    if paypal_client_id:
        config['paypal_client_id'] = paypal_client_id

    if not config:
        return jsonify({'error': 'Payment configuration not available'}), 500

    return jsonify(config), 200

@app.route('/api/subscription/upgrade', methods=['POST'])
@token_required
def upgrade_subscription():
    """Upgrade user to premium subscription."""
    try:
        data = request.get_json()
        plan_type = data.get('plan_type', 'monthly')  # monthly or yearly
        gateway = data.get('gateway', 'paystack')  # paystack or paypal

        # Define pricing
        prices = {
            'monthly': {'paystack': 1500000, 'paypal': 1500},  # 15000 NGN for Paystack, 15 USD for PayPal
            'yearly': {'paystack': 18000000, 'paypal': 18000}   # 180000 NGN for Paystack, 180 USD for PayPal
        }

        if plan_type not in prices or gateway not in ['paystack', 'paypal']:
            return jsonify({'error': 'Invalid plan type or payment gateway'}), 400

        amount = prices[plan_type][gateway]

        # Get user email
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT email FROM users WHERE id = ?', (g.user_id,))
        user = cursor.fetchone()
        if not user:
            return jsonify({'error': 'User not found'}), 404

        # Use the same payment initialization logic
        # Get user email
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT email FROM users WHERE id = ?', (g.user_id,))
        user = cursor.fetchone()
        if not user:
            return jsonify({'error': 'User not found'}), 404

        # Initialize payment using the same logic as initialize_payment
        if gateway == 'paypal':
            # PayPal integration
            if not PAYPAL_CLIENT_ID or not PAYPAL_CLIENT_SECRET:
                return jsonify({'error': 'PayPal payment service not configured'}), 500

            access_token = get_paypal_access_token()
            if not access_token:
                return jsonify({'error': 'Failed to authenticate with PayPal'}), 500

            # Convert amount to USD cents format for PayPal
            amount_usd = amount / 100  # Assuming amount comes in cents

            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {access_token}'
            }

            paypal_payload = {
                'intent': 'CAPTURE',
                'purchase_units': [{
                    'amount': {
                        'currency_code': 'USD',
                        'value': f'{amount_usd:.2f}'
                    },
                    'description': f'ATS Matcher Subscription Upgrade - {plan_type.title()}'
                }],
                'application_context': {
                    'return_url': PAYSTACK_CALLBACK_URL,
                    'cancel_url': f'{PAYSTACK_CALLBACK_URL}?payment=cancelled'
                }
            }

            response = requests.post(f'{PAYPAL_BASE_URL}/v2/checkout/orders', headers=headers, json=paypal_payload)
            result = response.json()

            if response.status_code == 201:
                # Find approval URL
                approval_url = None
                for link in result.get('links', []):
                    if link.get('rel') == 'approve':
                        approval_url = link['href']
                        break

                return jsonify({
                    'status': True,
                    'data': {
                        'authorization_url': approval_url,
                        'reference': result['id'],
                        'gateway': 'paypal'
                    }
                }), 200
            else:
                return jsonify({'error': 'PayPal payment initialization failed', 'details': result}), 400

        else:
            # Paystack integration (default)
            if not PAYSTACK_SECRET_KEY:
                return jsonify({'error': 'Paystack payment service not configured'}), 500

            url = f"{PAYSTACK_BASE_URL}/transaction/initialize"
            headers = {
                'Authorization': f'Bearer {PAYSTACK_SECRET_KEY}',
                'Content-Type': 'application/json'
            }
            payload = {
                'email': user['email'],
                'amount': str(amount),
                'callback_url': PAYSTACK_CALLBACK_URL,
                'metadata': {
                    'user_id': g.user_id,
                    'type': 'subscription_upgrade',
                    'plan_type': plan_type,
                    'gateway': gateway
                }
            }

            response = requests.post(url, headers=headers, json=payload)
            result = response.json()

            if response.status_code == 200 and result.get('status'):
                return jsonify(result), 200
            else:
                return jsonify({'error': 'Payment initialization failed', 'details': result}), 400

    except Exception as e:
        print(f"Subscription upgrade error: {e}")
        return jsonify({'error': 'Subscription upgrade failed'}), 500


def generate_standard_resume_pdf(resume_text):
    """
    Creates a clean, professional PDF resume by parsing the raw text and formatting it into standard CV sections.
    """
    
    # Use Gemini to parse the resume into structured sections
    parse_prompt = f"""
    Parse the following resume text into structured JSON format for a professional CV.
    
    Resume Text:
    {resume_text}
    
    Extract and structure the information into the following sections:
    - name: Full name of the person (string)
    - contact: Object with email, phone, location, linkedin (strings, use empty string if not found)
    - summary: Professional summary or objective (string, use empty string if not present)
    - experience: Array of work experience objects, each with: title (string), company (string), duration (string), location (string), description (array of strings)
    - education: Array of education objects, each with: degree (string), institution (string), year (string), gpa (string)
    - skills: Array of technical/professional skills (strings only)
    - certifications: Array of certification names (strings only, no JSON formatting)
    - projects: Array of project objects, each with: name (string), description (string or array of strings), technologies (array of strings)
    
    IMPORTANT: 
    - Return ONLY valid JSON. Do not add any other text.
    - All string values should be plain text without quotes, brackets, or JSON formatting.
    - Arrays should contain only the actual content strings.
    - If a section is not present, use empty array [] or empty string "".
    - Do not include any JSON-like formatting in the string values themselves.
    """
    
    try:
        response = model.generate_content(parse_prompt)
        response_text = getattr(response, 'text', None) or str(response)
        if not response_text:
            raise ValueError("Empty response from Gemini API")
        json_string = response_text.strip().replace("```json", "").replace("```", "").replace("```JSON", "")
        # Clean up any trailing/leading non-JSON text
        start_idx = json_string.find('{')
        end_idx = json_string.rfind('}') + 1
        if start_idx != -1 and end_idx > start_idx:
            json_string = json_string[start_idx:end_idx]
        if not json_string:
            raise ValueError("Empty JSON string after processing")
        parsed_data = json.loads(json_string)
        
        # Validate and clean the parsed data
        def clean_string(value):
            if isinstance(value, str):
                return value.strip().replace('{', '').replace('}', '').replace('"', '').replace("'", '')
            return str(value).strip()
        
        def clean_array(arr):
            if isinstance(arr, list):
                return [clean_string(item) for item in arr if clean_string(item)]
            elif isinstance(arr, str):
                # Handle case where array is returned as string
                return [item.strip() for item in arr.split(',') if item.strip()]
            return []
        
        # Clean each section
        parsed_data['name'] = clean_string(parsed_data.get('name', 'Professional Name'))
        parsed_data['summary'] = clean_string(parsed_data.get('summary', ''))
        parsed_data['skills'] = clean_array(parsed_data.get('skills', []))
        parsed_data['certifications'] = clean_array(parsed_data.get('certifications', []))
        
        # Clean contact info
        contact = parsed_data.get('contact', {})
        if isinstance(contact, dict):
            for key in contact:
                contact[key] = clean_string(contact[key])
        else:
            parsed_data['contact'] = {"email": "", "phone": "", "location": "", "linkedin": ""}
            
        # Clean experience
        if isinstance(parsed_data.get('experience'), list):
            for exp in parsed_data['experience']:
                if isinstance(exp, dict):
                    for key in exp:
                        if key == 'description' and isinstance(exp[key], list):
                            exp[key] = [clean_string(desc) for desc in exp[key]]
                        else:
                            exp[key] = clean_string(exp[key])
        
        # Clean projects
        if isinstance(parsed_data.get('projects'), list):
            for proj in parsed_data['projects']:
                if isinstance(proj, dict):
                    for key in proj:
                        if key == 'description':
                            if isinstance(proj[key], list):
                                proj[key] = [clean_string(desc) for desc in proj[key]]
                            else:
                                proj[key] = clean_string(proj[key])
                        elif key == 'technologies' and isinstance(proj[key], list):
                            proj[key] = [clean_string(tech) for tech in proj[key]]
                        else:
                            proj[key] = clean_string(proj[key])
        
    except Exception as e:
        print(f"Gemini API error: {e}")
        # Fallback: create basic structure
        parsed_data = {
            "name": "Professional Name",
            "contact": {"email": "", "phone": "", "location": "", "linkedin": ""},
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "certifications": [],
            "projects": []
        }
    
    # Generate PDF using ReportLab
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.units import inch
        from reportlab.lib import colors
    except Exception as e:
        raise RuntimeError(f"ReportLab not available: {e}")
    
    pdf_stream = BytesIO()
    doc = SimpleDocTemplate(pdf_stream, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(name='Name', fontSize=24, fontName='Helvetica-Bold', spaceAfter=20, alignment=1))
    styles.add(ParagraphStyle(name='Contact', fontSize=10, textColor=colors.gray, alignment=1, spaceAfter=20))
    styles.add(ParagraphStyle(name='SectionHeader', fontSize=14, fontName='Helvetica-Bold', textColor=colors.darkblue, spaceAfter=10, borderWidth=1, borderColor=colors.lightgrey, borderPadding=5))
    styles.add(ParagraphStyle(name='JobTitle', fontSize=12, fontName='Helvetica-Bold', textColor=colors.darkblue))
    styles.add(ParagraphStyle(name='JobCompany', fontSize=11, fontName='Helvetica-Oblique', textColor=colors.grey))
    styles.add(ParagraphStyle(name='JobDuration', fontSize=10, textColor=colors.grey, spaceAfter=10))
    styles.add(ParagraphStyle(name='NormalIndented', fontSize=10, leftIndent=20))
    
    story = []
    
    # Header
    name = parsed_data.get('name') or 'Professional Name'
    story.append(Paragraph(str(name), styles['Name']))
    
    contact_parts = []
    contact = parsed_data.get('contact', {}) or {}
    if contact.get('email'): contact_parts.append(str(contact['email']))
    if contact.get('phone'): contact_parts.append(str(contact['phone']))
    if contact.get('location'): contact_parts.append(str(contact['location']))
    if contact.get('linkedin'): contact_parts.append(str(contact['linkedin']))
    
    story.append(Paragraph(' | '.join(contact_parts), styles['Contact']))
    
    # Summary
    summary = parsed_data.get('summary')
    if summary:
        story.append(Paragraph('PROFESSIONAL SUMMARY', styles['SectionHeader']))
        story.append(Paragraph(str(summary), styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Experience
    if parsed_data.get('experience'):
        story.append(Paragraph('PROFESSIONAL EXPERIENCE', styles['SectionHeader']))
        for exp in parsed_data['experience']:
            title = exp.get('title') or ''
            story.append(Paragraph(str(title), styles['JobTitle']))
            company_info = str(exp.get('company', ''))
            if exp.get('location'):
                company_info += f" - {str(exp.get('location', ''))}"
            story.append(Paragraph(company_info, styles['JobCompany']))
            if exp.get('duration'):
                story.append(Paragraph(str(exp.get('duration', '')), styles['JobDuration']))
            
            if exp.get('description'):
                for desc in exp['description']:
                    story.append(Paragraph(f"• {str(desc)}", styles['NormalIndented']))
            story.append(Spacer(1, 12))
    
    # Education
    if parsed_data.get('education'):
        story.append(Paragraph('EDUCATION', styles['SectionHeader']))
        for edu in parsed_data['education']:
            degree = str(edu.get('degree', ''))
            institution = str(edu.get('institution', ''))
            year = str(edu.get('year', ''))
            gpa = str(edu.get('gpa', ''))
            
            edu_text = f"<b>{degree}</b><br/>{institution}"
            if year:
                edu_text += f", {year}"
            if gpa:
                edu_text += f" (GPA: {gpa})"
            
            story.append(Paragraph(edu_text, styles['Normal']))
            story.append(Spacer(1, 6))
    
    # Skills
    if parsed_data.get('skills'):
        story.append(Paragraph('SKILLS', styles['SectionHeader']))
        skills_text = ', '.join(str(skill).strip() for skill in parsed_data['skills'] if str(skill).strip())
        story.append(Paragraph(skills_text, styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Certifications
    if parsed_data.get('certifications'):
        story.append(Paragraph('CERTIFICATIONS', styles['SectionHeader']))
        for cert in parsed_data['certifications']:
            cert_text = str(cert).strip()
            # Clean up any JSON formatting artifacts
            cert_text = cert_text.replace('{', '').replace('}', '').replace('"', '').replace("'", '').strip()
            if cert_text:
                story.append(Paragraph(f"• {cert_text}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Projects
    if parsed_data.get('projects'):
        story.append(Paragraph('PROJECTS', styles['SectionHeader']))
        for proj in parsed_data['projects']:
            name = str(proj.get('name', '')).strip()
            if name:
                story.append(Paragraph(name, styles['JobTitle']))
            
            description = proj.get('description')
            if description:
                if isinstance(description, list):
                    for desc_item in description:
                        desc_text = str(desc_item).strip()
                        if desc_text:
                            story.append(Paragraph(f"• {desc_text}", styles['NormalIndented']))
                else:
                    desc_text = str(description).strip()
                    if desc_text:
                        story.append(Paragraph(desc_text, styles['Normal']))
            
            technologies = proj.get('technologies')
            if technologies:
                if isinstance(technologies, list):
                    tech_text = f"Technologies: {', '.join(str(tech).strip() for tech in technologies if str(tech).strip())}"
                else:
                    tech_text = f"Technologies: {str(technologies).strip()}"
                story.append(Paragraph(tech_text, styles['NormalIndented']))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    pdf_stream.seek(0)
    return pdf_stream


def generate_optimized_resume_docx(original_text, missing_skills, analysis_data=None):
    """Create a comprehensive .docx document with full analysis data."""
    doc = Document()
    doc.add_heading('Optimized Resume Draft (AI Suggestions)', level=1)

    doc.add_heading('AI Recommended Skill Enhancement Section', level=2)
    doc.add_paragraph('Based on the job requirements, consider integrating these keywords into your experience or skill sections:')
    for skill in missing_skills:
        doc.add_paragraph(f'• {skill} (Targeted Keyword)', style='List Bullet')

    # Add Gap Analysis if available
    if analysis_data and analysis_data.get("keyword_gap_analysis"):
        doc.add_heading('Keyword Gap Analyzer', level=2)
        for skill, section in analysis_data["keyword_gap_analysis"].items():
            doc.add_paragraph(f'{skill} → Add to: {section}', style='List Bullet')

    # Add Weakly Represented Skills if available
    if analysis_data and analysis_data.get("weakly_represented_skills"):
        doc.add_heading('Weakly Represented Skills (Needs More Emphasis)', level=2)
        for skill in analysis_data["weakly_represented_skills"]:
            doc.add_paragraph(f'• {skill}', style='List Bullet')

    # Add Overused Terms if available
    if analysis_data and analysis_data.get("overused_terms"):
        doc.add_heading('Overused Terms (Consider Varying)', level=2)
        for term in analysis_data["overused_terms"]:
            doc.add_paragraph(f'• {term}', style='List Bullet')

    # Add Add to Resume Suggestions if available
    if analysis_data and analysis_data.get("add_to_resume_suggestions"):
        doc.add_heading('Add to Resume Suggestions', level=2)
        for suggestion in analysis_data["add_to_resume_suggestions"]:
            doc.add_paragraph(f'• {suggestion}', style='List Bullet')

    doc.add_heading('Extracted Original Resume Content', level=2)
    # Add the original resume text as a single paragraph (preserve basic newlines)
    for line in original_text.splitlines():
        if line.strip():
            doc.add_paragraph(line)

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# Helper function to extract text from PDF
def extract_text_from_pdf(pdf_stream):
    """Reads PDF file stream and extracts all text content."""
    try:
        # PyPDF2 requires the stream object
        reader = PyPDF2.PdfReader(pdf_stream)
        text = ""
        for page in reader.pages:
            # We use .extract_text() to get the text content
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return None

@app.route('/api/match', methods=['POST'])
@token_required
def process_match():
    """API endpoint to receive resume and JD, process with Gemini, and return structured JSON plus extracted text.."""

    # Check usage limits
    can_use, usage_message = check_usage_limit(g.user_id, 'analysis')
    if not can_use:
        return jsonify({
            "error": usage_message,
            "upgrade_required": True
        }), 429  # Too Many Requests

    if not model:
        return jsonify({"error": "Gemini API model is not available. Check server logs."}), 500

    # Ensure all required data is present
    if 'resume' not in request.files or 'job_description' not in request.form:
        return jsonify({"error": "Missing resume file or job description in the request."}), 400

    # Extract data from the POST request
    resume_file = request.files['resume']
    job_description = request.form['job_description']

    # 1. Convert PDF to Text
    # We pass the stream object from the request file to our helper function
    resume_text = extract_text_from_pdf(resume_file.stream)

    if not resume_text or len(resume_text.strip()) < 50:
        return jsonify({"error": "Could not extract sufficient text from the PDF file. Is the PDF text-searchable?"}), 400

    # 2. Construct the Prompt for Gemini
    prompt = f"""
    You are an expert Applicant Tracking System (ATS) Analyst and Resume Optimization Specialist. Your job is to compare a RESUME against a JOB DESCRIPTION with comprehensive scoring and recommendations.

    --- RESUME TEXT ---
    {resume_text}

    --- JOB DESCRIPTION TEXT ---
    {job_description}

    Analyze the two texts and provide a comprehensive evaluation:

    1. **Calculate Detailed Scores (0-100%):**
       - "keyword_match_score": Percentage of critical job keywords present in resume
       - "skills_alignment_score": How well candidate's skills align with job requirements
       - "experience_relevance_score": How relevant work experience is to the position
       - "formatting_structure_score": How well-structured resume matches ATS expectations
       - "seniority_fit_score": Whether experience level matches position seniority
       - "overall_match_score": Weighted average of all scores

    2. **Identify Matched Skills:** List 5-10 key professional skills/technologies present in BOTH documents.

    3. **Identify Missing Skills:** List 5-10 key professional skills/technologies required by JOB DESCRIPTION but NOT found in RESUME.

    4. **Identify Weakly Represented Skills:** List 3-5 skills that appear in both documents but with weak representation in the resume (mentioned once or briefly).

    5. **Identify Overused Terms:** List any keywords/phrases that appear excessively in the resume (3+ times) that should be varied.

    6. **Keyword Gap Analysis:** For missing skills, suggest specific resume sections where each could be naturally integrated.

    7. **Generate Recommendation:** Write brief (max 3 sentences), actionable advice to improve resume for this specific job.

    8. **Add to Resume Suggestions:** Provide 3-5 specific bullet points or phrases the candidate could add to their resume to improve match.

    Return results STRICTLY as a single JSON object. Do not add any other text.
    Required keys: "keyword_match_score", "skills_alignment_score", "experience_relevance_score", "formatting_structure_score", "seniority_fit_score", "overall_match_score", "matched_skills", "missing_skills", "weakly_represented_skills", "overused_terms", "keyword_gap_analysis", "recommendation_text", "add_to_resume_suggestions"

    All scores must be integers 0-100.
    All list fields must be arrays of strings.
    keyword_gap_analysis must be an object mapping missing skill to suggested resume section (e.g., {{"Python": "Technical Skills section", "Docker": "Projects section"}}).
    """

    try:
        # 3. Call the Gemini API
        print("Sending prompt to Gemini API via GenerativeModel...")
        response = model.generate_content(prompt)

        # 4. Parse the JSON response
        # We use a robust method to extract the JSON from the model's text output
        json_string = getattr(response, 'text', str(response)).strip().replace("```json", "").replace("```", "")
        result_json = json.loads(json_string)

        # Add backward compatibility: if 'overall_match_score' exists, also set 'score'
        if 'overall_match_score' in result_json and 'score' not in result_json:
            result_json['score'] = result_json['overall_match_score']

        print(f"Analysis successful. Overall Score: {result_json.get('overall_match_score', result_json.get('score', 'N/A'))}")

        # Record usage after successful analysis
        record_usage(g.user_id, 'analysis')

        result_json['original_resume_text'] = resume_text  # Include original text for later use
        return jsonify(result_json)

    except json.JSONDecodeError:
        print(f"ERROR: Failed to parse JSON from Gemini response: {response.text}")
        return jsonify({"error": "Analysis failed to produce valid JSON output."}), 500
    except Exception as e:
        print(f"Gemini API error: {e}")
        return jsonify({"error": f"Internal analysis failed due to API error: {e}"}), 500




@app.route('/api/generate-cv', methods=['POST'])
@token_required
def generate_cv():
    """Endpoint to generate and return the optimized PDF file with full analysis."""
    
    data = request.json
    original_resume_text = data.get('original_resume_text')
    missing_skills = data.get('missing_skills', [])
    
    # Extract full analysis data if provided
    analysis_data = {
        'keyword_gap_analysis': data.get('keyword_gap_analysis'),
        'weakly_represented_skills': data.get('weakly_represented_skills'),
        'overused_terms': data.get('overused_terms'),
        'add_to_resume_suggestions': data.get('add_to_resume_suggestions')
    }
    
    if not original_resume_text or not missing_skills:
        return jsonify({"error": "Missing original resume text or missing skills data."}), 400
        
    # Generate the PDF stream (handle missing WeasyPrint gracefully)
    # Try PDF first, but fallback to DOCX if WeasyPrint or system deps are missing
    try:
        pdf_stream = generate_optimized_resume_pdf(original_resume_text, missing_skills, analysis_data)
        return send_file(
            pdf_stream,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='optimized_cv.pdf'
        )
    except RuntimeError as e:
        # RuntimeError from the PDF generator indicates WeasyPrint or its deps missing.
        try:
            docx_stream = generate_optimized_resume_docx(original_resume_text, missing_skills, analysis_data)
            return send_file(
                docx_stream,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='optimized_cv.docx'
            )
        except Exception as ex:
            return jsonify({"error": f"Failed to generate fallback DOCX: {ex}"}), 500
    except Exception as e:
        return jsonify({"error": f"Failed to generate PDF: {e}"}), 500


@app.route('/api/generate-standard-resume', methods=['POST'])
@token_required
def generate_standard_resume():
    """Endpoint to generate and return a clean, standard PDF resume."""
    
    try:
        data = request.get_json()
    except Exception as e:
        return jsonify({"error": "Invalid JSON in request."}), 400
    
    if not data:
        return jsonify({"error": "No JSON data in request."}), 400
        
    resume_text = data.get('resume_text')
    
    if not resume_text:
        return jsonify({"error": "Missing resume text."}), 400
    
    if not model:
        return jsonify({"error": "Gemini API model is not available. Check server logs."}), 500
    
    try:
        pdf_stream = generate_standard_resume_pdf(resume_text)
        return send_file(
            pdf_stream,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='standard_resume.pdf'
        )
    except RuntimeError as e:
        return jsonify({"error": f"PDF generation failed: {e}"}), 500
    except Exception as e:
        return jsonify({"error": f"Failed to generate standard resume: {e}"}), 500


if __name__ == '__main__':
    # Initialize the database
    init_db()

    # Ensure you set GEMINI_API_KEY environment variable first!
    # To run: python app.py
    # This will run the backend server on port 5000
    app.run(debug=False, port=5000)